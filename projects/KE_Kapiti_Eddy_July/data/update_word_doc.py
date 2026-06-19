import openpyxl
import docx
import pandas as pd
import re

xlsm_path = 'Species_params_GEMINI.xlsm'
docx_path = 'Species_params_word_table_updated.docx'
out_docx_path = 'Species_params_word_table_Final.docx'

print(f"Loading XLSM data from {xlsm_path}...")
# Use data_only=False so we can read the raw formula strings (e.g. "=J2")
wb = openpyxl.load_workbook(xlsm_path)
ws = wb.active

header_row = [cell.value for cell in ws[1]]
header_idx = {val: idx for idx, val in enumerate(header_row, start=1) if val is not None}

sci_col = header_idx['scientific_name']
code_col = header_idx['code']
midpoint_col = header_idx['midpoint']
ref_col = header_idx['references']

def resolve_value(ws, row, col):
    val = ws.cell(row=row, column=col).value
    visited = set()
    while isinstance(val, str) and val.startswith('='):
        ref_cell = val[1:]
        if ref_cell in visited:
            break
        visited.add(ref_cell)
        try:
            val = ws[ref_cell].value
        except:
            break
    return val

lookup = {}
for r in range(2, ws.max_row + 1):
    sci_val = resolve_value(ws, r, sci_col)
    code_val = resolve_value(ws, r, code_col)
    
    if sci_val is None or code_val is None:
        continue
        
    sci = str(sci_val).strip().lower()
    code = str(code_val).strip().upper()
    
    if sci == 'none' or code == 'none': continue
    
    # Normalize species names to simple keys for mapping
    if 'cynodon' in sci: sci = 'cynodon'
    elif 'cenchrus' in sci: sci = 'cenchrus'
    elif 'tephrosia' in sci: sci = 'tephrosia'
    elif 'vachellia' in sci or 'drepallobium' in sci: sci = 'vachellia'
    elif 'themeda' in sci or 'triandra' in sci: sci = 'themeda'
    elif 'indigofera' in sci or 'volkensii' in sci: sci = 'indigofera'
    
    midpoint = resolve_value(ws, r, midpoint_col)
    references = resolve_value(ws, r, ref_col)
    
    if pd.notna(midpoint):
        lookup[(sci, code)] = {
            'midpoint': midpoint,
            'references': str(references).strip() if pd.notna(references) and str(references).strip() != 'None' else ''
        }

print(f"Loaded {len(lookup)} valid parameters from XLSM.")

print(f"Loading Word document {docx_path}...")
doc = docx.Document(docx_path)

ref_list = []
ref_map = {} # map string -> index (1-based)

def get_ref_index(ref_str):
    if not ref_str:
        return None
    if ref_str not in ref_map:
        ref_list.append(ref_str)
        ref_map[ref_str] = len(ref_list)
    return ref_map[ref_str]

def extract_code(desc):
    # Fix spacing issues in Word doc (e.g. H2OREF_ GS -> H2OREF_GS)
    desc = desc.replace('H2OREF_ GS', 'H2OREF_GS')
    desc = desc.replace('H2OREF_LEAFGROWTH', 'H2OREF_LEAF_GROWTH')
    matches = re.findall(r'\b([A-Z0-9_]{3,})\b', desc)
    known_codes = [
        'AEJM', 'AEKC', 'AEKO', 'AERD', 'AEVC', 'AEVO', 'AEV0',
        'KC25', 'KM20', 'VCMAX', 'VCMAX25', 'THETA',
        'GSMAX', 'GSMIN', 'WUECMAX', 'WUECMIN',
        'H2OREF_A', 'H2OREF_SENESCENCE', 'H2OREF_GS',
        'H2OREF_FLUSHING', 'H2OREF_LEAF_GROWTH',
        'GDD_BASE_TEMPERATURE', 'GDD_EMERG', 'GDD_EMERGENCE',
        'GDD_MATURITY', 'GDD_STEM_ELONGATION',
        'GDD_ROOTS_GROWN', 'GDDFOLSTART', 'GDDFOLEND',
        'NDFLUSH', 'NDMORTA', 'FRACTION_ROOT',
        'FRACTION_FOLIAGE', 'FRACTION_FRUIT',
        'MFOLOPT', 'MWFM', 'SLAMAX', 'SLAMIN', 'KO25',
        'QJVC', 'QRD25', 'AMAXB', 'AMX25'
    ]
    for match in reversed(matches):
        if match in known_codes:
            return match
    return None

def format_val(val):
    try:
        fval = float(val)
        if fval == int(fval):
            return str(int(fval))
        return f"{fval:.6g}"
    except:
        return str(val)

# Process Table 1
table1 = doc.tables[0]
t1_species_row = table1.rows[2]
t1_col_map = {}
for j, cell in enumerate(t1_species_row.cells[2:], start=2):
    txt = cell.text.lower()
    if 'nlemfuensis' in txt or 'cynodon' in txt: t1_col_map[j] = 'cynodon'
    elif 'mezianum' in txt or 'cenchrus' in txt: t1_col_map[j] = 'cenchrus'
    elif 'triandra' in txt or 'themeda' in txt: t1_col_map[j] = 'themeda'
    elif 'volkensii' in txt or 'indigofera' in txt: t1_col_map[j] = 'indigofera'
    elif 'drepallobium' in txt or 'vachellia' in txt: t1_col_map[j] = 'vachellia'

for i in range(3, len(table1.rows)):
    row = table1.rows[i]
    desc = row.cells[0].text
    code = extract_code(desc)
    if not code: continue
    
    # Handle synonyms in the XLSM
    if code == 'VCMAX25': code = 'VCMAX'
    if code == 'AEV0': code = 'AEVO'
    
    for j, sci_key in t1_col_map.items():
        lookup_code = code
        key = (sci_key, lookup_code)
        
        if key not in lookup and lookup_code == 'VCMAX':
            key = (sci_key, 'VCMAX25')
        
        if key in lookup:
            data = lookup[key]
            val_str = format_val(data['midpoint'])
            ref_idx = get_ref_index(data['references'])
            
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.add_run(val_str)
            if ref_idx is not None:
                p.add_run(f'[{ref_idx}]').font.superscript = True

# Process Table 2
if len(doc.tables) > 1:
    table2 = doc.tables[1]
    t2_species_row = table2.rows[1]
    t2_col_map = {}
    for j, cell in enumerate(t2_species_row.cells[2:], start=2):
        txt = cell.text.lower()
        if 'dactylon' in txt or 'cynodon' in txt: t2_col_map[j] = 'cynodon'
        elif 'mezianus' in txt or 'cenchrus' in txt: t2_col_map[j] = 'cenchrus'
        elif 'triandra' in txt or 'themeda' in txt: t2_col_map[j] = 'themeda'
        elif 'volkensii' in txt or 'indigofera' in txt: t2_col_map[j] = 'indigofera'
        elif 'drepallobium' in txt or 'vachellia' in txt: t2_col_map[j] = 'vachellia'

    for i in range(2, len(table2.rows)):
        row = table2.rows[i]
        desc = row.cells[0].text
        code = extract_code(desc)
        if not code: continue
        
        if code == 'VCMAX25': code = 'VCMAX'
        if code == 'AEV0': code = 'AEVO'
        
        for j, sci_key in t2_col_map.items():
            lookup_code = code
            key = (sci_key, lookup_code)
            
            if key not in lookup and lookup_code == 'VCMAX':
                key = (sci_key, 'VCMAX25')
            
            if key in lookup:
                data = lookup[key]
                val_str = format_val(data['midpoint'])
                ref_idx = get_ref_index(data['references'])
                
                cell = row.cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                p.add_run(val_str)
                if ref_idx is not None:
                    p.add_run(f'[{ref_idx}]').font.superscript = True

# Add References Section
doc.add_page_break()
doc.add_heading('References', level=1)
for i, ref in enumerate(ref_list, start=1):
    doc.add_paragraph(f"[{i}] {ref}")

doc.save(out_docx_path)
print(f"Successfully generated {out_docx_path}")
